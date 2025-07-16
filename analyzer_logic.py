# analyzer_logic.py

import os
import json
import fitz  # PyMuPDF
import io
import re
import google.generativeai as genai
from dotenv import load_dotenv
import docx
from json.decoder import JSONDecodeError
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random
from docx.enum.table import WD_TABLE_ALIGNMENT , WD_CELL_VERTICAL_ALIGNMENT

load_dotenv()

# --- Gemini Client Initialization ---
model = None
try:
    gemini_key = os.environ.get("GEMINI_API_KEY")
    if not gemini_key:
        print("FATAL ERROR: GEMINI_API_KEY environment variable not found.")
    else:
        genai.configure(api_key=gemini_key)
        model = genai.GenerativeModel('gemini-1.5-pro')
        print("Gemini model 'gemini-1.5-flash' initialized successfully.")
except Exception as e:
    print(f"Error initializing Gemini client: {e}")

def extract_text_from_docx_stream(docx_stream):
    try:
        document = docx.Document(docx_stream)
        return "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX stream: {e}")
        return None

def extract_text_from_pdf_stream(pdf_stream):
    try:
        doc = fitz.open(stream=pdf_stream, filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
        return text
    except Exception as e:
        print(f"Error reading PDF stream: {e}")
        return None

def extract_text_from_file(file_storage):
    filename = file_storage.filename.lower()
    file_stream = io.BytesIO(file_storage.read())
    if filename.endswith('.pdf'):
        return extract_text_from_pdf_stream(file_stream)
    elif filename.endswith('.docx'):
        return extract_text_from_docx_stream(file_stream)
    else:
        print(f"Unsupported file format: {filename}")
        return None

def analyze_resume_with_ai(resume_text, jd_text, initial_analysis=None):
    if not model:
        return {"error": "AI client not initialized. Check server logs for API Key issues."}

    generation_config = {
      "temperature": 0.2,
      "response_mime_type": "application/json",
    }
    
    # Gemini works best by combining system instructions with the user query
    # into a single, comprehensive prompt.
    if initial_analysis:
        full_prompt = (
            "You are an expert ATS Re-Analyzer. Your task is to evaluate an UPDATED resume against a job description, "
            "specifically assessing how well it incorporated previous feedback. Your new score MUST reflect the improvements made. "
            "Recognize when keywords and suggestions have been successfully integrated and score more generously than a first-pass analysis. "
            "Your entire output MUST be a single, valid JSON object.\n\n"
            "Your task is to re-evaluate the resume. You are given the initial analysis which pointed out gaps. Now, review the new resume text and assess the improvements.\n\n"
            f"--- JOB DESCRIPTION ---\n{jd_text}\n\n"
            f"--- INITIAL ANALYSIS & SUGGESTIONS ---\n{json.dumps(initial_analysis, indent=2)}\n\n"
            f"--- NEW, UPDATED RESUME TEXT ---\n{resume_text}\n\n"
            "--- RE-ANALYSIS REQUIREMENTS ---\n"
            "Generate a new JSON analysis. The scores in this new analysis should be HIGHER than the initial analysis if the suggestions were followed. "
            "Use the same JSON structure as before:\n"
            """
            {
                "summary": "[2-3 sentences explaining how the resume has improved and any remaining gaps]",
                "strengths": ["List the strongest matching skills/experiences in the NEW resume"],
                "missing_keywords": ["List any CRITICAL keywords that are still missing, if any"],
                "suggested_changes": ["Provide 1-2 final polish suggestions if needed, otherwise an empty list"],
                "scoring_breakdown": {
                    "key_skills": { "score": "[Integer 0-100]", "justification": "[Justify score based on improvement]" },
                    "experience_level": { "score": "[Integer 0-100]", "justification": "[Justify score based on improvement]" },
                    "project_and_impact": { "score": "[Integer 0-100]", "justification": "[Justify score based on improvement]" },
                    "education_and_certs": { "score": "[Integer 0-100]", "justification": "[Justify score based on improvement]" }
                }
            }
            """
        )
    else:
        full_prompt = (
            "You are an expert ATS (Applicant Tracking System) analyzer with 15 years of experience in recruitment. "
            "Your role is to provide a brutally honest, realistic assessment of resume-job fit by breaking down the analysis into specific, scored categories. "
            "You must be STRICT in your scoring for each category. Do not inflate scores. "
            "Your entire output MUST be a single, valid JSON object, and nothing else. Do not wrap the JSON in markdown code blocks.\n\n"
            f"Analyze this resume against the job description using STRICT ATS criteria. Be realistic and critical in your assessment.\n\n"
            f"--- JOB DESCRIPTION ---\n{jd_text}\n\n"
            f"--- RESUME TEXT ---\n{resume_text}\n\n"
            "--- ANALYSIS REQUIREMENTS ---\n"
            "Generate a JSON object with this exact structure:\n"
            """
            {
                "summary": "[2-3 sentences explaining the fit level and main gaps]",
                "strengths": ["List actual skills/experiences that STRONGLY match the JD"],
                "missing_keywords": ["Critical keywords/skills completely absent from resume"],
                "suggested_changes": ["Specific, actionable commands for improving the resume"],
                "scoring_breakdown": {
                    "key_skills": { "score": "[Integer 0-100]", "justification": "[Justify score]" },
                    "experience_level": { "score": "[Integer 0-100]", "justification": "[Justify score]" },
                    "project_and_impact": { "score": "[Integer 0-100]", "justification": "[Justify score]" },
                    "education_and_certs": { "score": "[Integer 0-100]", "justification": "[Justify score]" }
                }
            }
            """
        )

    try:
        response = model.generate_content(full_prompt, generation_config=generation_config)
        
        # Accessing response content is different in Gemini
        if not response.parts:
            return {"error": "Request failed or was filtered by the AI."}
        
        result = json.loads(response.text)
        
        breakdown = result.get("scoring_breakdown", {})
        weights = {
            "key_skills": 0.40, "experience_level": 0.30,
            "project_and_impact": 0.20, "education_and_certs": 0.10
        }
        
        total_score = 0
        if all(k in breakdown and isinstance(breakdown[k], dict) and "score" in breakdown[k] for k in weights.keys()):
            for category, weight in weights.items():
                score_str = str(breakdown[category].get("score", "0")).strip('%')
                score = int(score_str) if score_str.isdigit() else 0
                total_score += score * weight
            
            total_score += random.randint(-2, 2)
            calculated_score = min(98, max(5, int(total_score)))
        else:
            calculated_score = 42 
            result["summary"] = "AI response for scoring was malformed. This is a fallback score. " + result.get("summary", "")

        result["match_score"] = calculated_score
        
        return result

    except JSONDecodeError:
        return {"error": "The AI returned a response in an invalid format. Please try again."}
    except Exception as e:
        print(f"An unexpected server-side AI error occurred in analyze_resume_with_ai: {type(e).__name__} - {e}")
        return {"error": f"An unexpected server-side AI error occurred. Please check the server logs."}


def generate_new_resume_text_with_ai(original_resume_text, jd_text, suggested_changes, reformat_only=False, candidate_name="", job_title_only=""):
    if not model:
        return {"error": "AI client not initialized. Check server logs for API Key issues."}
    if not candidate_name:
        return {"error": "Candidate name was not provided to the generation function."}

    generation_config = {
      "temperature": 0.3,
      "response_mime_type": "application/json",
    }

    json_structure_prompt = f"""
--- OUTPUT FORMAT (NON-NEGOTIABLE) ---
Your entire output MUST be a single, valid JSON object. Do not add any other text or markdown.
The JSON object must have this top-level structure:
{{
    "candidate_name": "{candidate_name}",
    "designation_line": "[string] e.g., 'Senior Software Engineer | 10+ Years of Experience'",
    "contact_info": {{ "phone": "[string]", "email": "[string]" }},
    "sections": [
        // Array of section objects. See examples below.
    ]
}}

--- SECTION FORMATS ---
You must use one of the following formats for each object inside the "sections" array:

1.  *Standard Section (for Summary, Skills, Education, etc.):*
    {{
        "title": "[string] The section title",
        "content": "[string or array of strings] For paragraphs, a single string. For lists, an array of strings."
    }}

2.  *Experience Section (USE ONLY FOR A SECTION TITLED 'Experience'):*
    Its "title" MUST be "Experience" and its "content" MUST be an array of job objects:
    {{
        "title": "Experience",
        "content": [
            {{
                "job_title": "Senior AI Engineer",
                "company_and_date": "Innovate Corp | Jan 2020 - Present",
                "duties": [
                    "Led the development of a real-time sentiment analysis engine...",
                    "Quantified achievement like: 'Improved model accuracy by 15%...'"
                ]
            }}
        ]
    }}

3.  *Projects Section (USE ONLY FOR A SECTION TITLED 'Projects'):*
    Its "title" MUST be "Projects" and its "content" MUST be an array of project objects:
    {{
        "title": "Projects",
        "content": [
            {{
                "project_name": "Project Title Here",
                "description": "A detailed description of the project, including key actions and achievements.",
                "tech_stack": "List, of, technologies, used"
            }}
        ]
    }}
"""

    if reformat_only:
        system_prompt = "You are a resume data structuring robot."
        user_prompt = f"""
        Parse the following resume text into the required JSON format. Preserve the original content exactly.
        --- ORIGINAL RESUME ---
        {original_resume_text}
        {json_structure_prompt}
        """
    elif job_title_only:
        system_prompt = "You are an expert resume writer and career coach."
        user_prompt = f"""
        Your mission: Rewrite the original resume to be perfectly tailored for the job title: "{job_title_only}".
        Do not invent new experiences, but rephrase existing duties and projects to highlight skills relevant to this role.
        Quantify achievements where possible. Create a compelling summary that aligns with the target title.
        
        --- ORIGINAL RESUME ---
        {original_resume_text}
        
        {json_structure_prompt}
        
        Now, generate the complete, rewritten resume as a single JSON object targeted for a '{job_title_only}' position.
        """
    else:
        system_prompt = "You are an expert resume writer with deep ATS knowledge."
        user_prompt = f"""
        Your mission: Transform the original resume to achieve an 85%+ match score with the target job description, and output the result as a structured JSON object.
        Integrate missing keywords naturally into the project 'description' or experience bullet points. Quantify achievements.
        --- TARGET JOB DESCRIPTION ---
        {jd_text}
        --- ORIGINAL RESUME ---
        {original_resume_text}
        --- AI ANALYSIS & SUGGESTIONS ---
        {json.dumps(suggested_changes, indent=2)}
        {json_structure_prompt}
        Now, generate the complete, rewritten resume as a single JSON object.
        """
    
    full_prompt = f"{system_prompt}\n\n{user_prompt}"

    try:
        response = model.generate_content(full_prompt, generation_config=generation_config)
        
        if not response.parts:
            return {"error": "Rewrite failed or was filtered by the AI."}
        
        result_json = json.loads(response.text)
        return {"new_resume_json": result_json}
    
    except JSONDecodeError as e:
        return {"error": "The AI returned a response in an invalid JSON format."}
    except Exception as e:
        return {"error": f"An unexpected server-side AI error occurred: {e}"}

def convert_resume_json_to_text(resume_json):
    if not isinstance(resume_json, dict): return ""
    parts = []
    parts.append(resume_json.get("candidate_name", ""))
    parts.append(resume_json.get("designation_line", ""))
    contact = resume_json.get("contact_info", {})
    contact_str = " | ".join(filter(None, [contact.get("email"), contact.get("phone")]))
    if contact_str: parts.append(contact_str)
    
    for section in resume_json.get("sections", []):
        parts.append("\n" + section.get("title", "").upper())
        content = section.get("content")
        title_lower = section.get("title", "").lower()

        if isinstance(content, list):
            if title_lower == "projects":
                for project in content:
                    if isinstance(project, dict):
                        project_parts = [f"Project: {project.get('project_name', 'N/A')}"]
                        if project.get('description'): project_parts.append(f"  Description: {project.get('description')}")
                        if project.get('tech_stack'): project_parts.append(f"  Tech Stack: {project.get('tech_stack')}")
                        parts.append("\n".join(filter(None, project_parts)))
            elif title_lower == "experience":
                for job in content:
                     if isinstance(job, dict):
                        job_parts = [f"{job.get('job_title', 'N/A')}"]
                        if job.get('company_and_date'): job_parts.append(job.get('company_and_date'))
                        if isinstance(job.get('duties'), list):
                            job_parts.extend([f"- {duty}" for duty in job.get('duties')])
                        parts.append("\n".join(filter(None, job_parts)))
            else: 
                parts.extend([f"- {item}" for item in content])
        elif isinstance(content, str): 
            parts.append(content)
            
    return "\n".join(filter(None, parts))

# --- PDF Generation Constants for better readability ---
TOP_MARGIN, BOTTOM_MARGIN, LEFT_MARGIN, RIGHT_MARGIN = 60, 60, 60, 60
FONT_NAME_REGULAR, FONT_NAME_BOLD, FONT_NAME_ITALIC = "helv", "helv", "helv"
COLOR_PRIMARY, COLOR_SECONDARY, COLOR_ACCENT, COLOR_LINE = (0.1, 0.1, 0.1), (0.3, 0.3, 0.3), (0.1, 0.3, 0.8), (0.8, 0.8, 0.8)
COLOR_TABLE_HEADER_BG = (0.93, 0.93, 0.93)
LINE_HEIGHT = 15.5
BODY_FONTSIZE = 11

def _get_wrapped_text_height(text, width, fontname=FONT_NAME_REGULAR, fontsize=BODY_FONTSIZE, line_height=LINE_HEIGHT):
    lines = 0
    words = str(text).split()
    if not words: return line_height
    current_line = ""
    for word in words:
        test_line = f"{current_line} {word}".strip()
        if fitz.get_text_length(test_line, fontname=fontname, fontsize=fontsize) < width:
            current_line = test_line
        else:
            lines += 1
            current_line = word
    lines += 1
    return lines * line_height

def _render_text_with_wrapping(page, y, text, width, x_offset=0, bullet=False, fontname=FONT_NAME_REGULAR, fontsize=BODY_FONTSIZE, color=COLOR_SECONDARY):
    x_pos = LEFT_MARGIN + x_offset
    text_width = width - x_offset
    
    if bullet:
        page.insert_text(fitz.Point(LEFT_MARGIN, y), "•", fontname=FONT_NAME_BOLD, fontsize=fontsize+2, color=COLOR_PRIMARY)
    
    words = str(text).strip().split()
    if not words: return y + LINE_HEIGHT

    current_line = ""
    for word in words:
        test_line = f"{current_line} {word}".strip()
        if fitz.get_text_length(test_line, fontname=fontname, fontsize=fontsize) < text_width:
            current_line = test_line
        else:
            page.insert_text(fitz.Point(x_pos, y), current_line, fontname=fontname, fontsize=fontsize, color=color)
            y += LINE_HEIGHT
            current_line = word
    page.insert_text(fitz.Point(x_pos, y), current_line, fontname=fontname, fontsize=fontsize, color=color)
    return y + LINE_HEIGHT

def create_pdf_with_logo(resume_data, company):
    logo_paths = { "beround": "static/logos/beround.jpg", "climber": "static/logos/climber.jpg", "rennova": "static/logos/rennova.jpg" }
    logo_path = logo_paths.get(company)
    if company in logo_paths and (not logo_path or not os.path.exists(logo_path)):
        raise FileNotFoundError(f"Logo for {company} not found at {logo_path}")
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = TOP_MARGIN
    page_num = 1
    if logo_path: page.insert_image(fitz.Rect(page.rect.width - RIGHT_MARGIN - 80, TOP_MARGIN - 40, page.rect.width - RIGHT_MARGIN, TOP_MARGIN - 10), filename=logo_path)
    candidate_name = resume_data.get("candidate_name", "Candidate Name")
    designation_line = resume_data.get("designation_line", "Professional")
    contact_info = resume_data.get("contact_info", {})
    contact_items = [item for item in [contact_info.get("phone"), contact_info.get("email")] if item]
    contact_y, contact_fontsize = TOP_MARGIN + 4, 9.5
    for item in contact_items:
        text_width = fitz.get_text_length(item, fontname=FONT_NAME_REGULAR, fontsize=contact_fontsize)
        page.insert_text(fitz.Point(page.rect.width - RIGHT_MARGIN - text_width, contact_y), item, fontname=FONT_NAME_REGULAR, fontsize=contact_fontsize, color=COLOR_SECONDARY)
        contact_y += 14
    
    page.insert_text(fitz.Point(LEFT_MARGIN, y), candidate_name, fontname=FONT_NAME_BOLD, fontsize=26, color=COLOR_PRIMARY)
    y += 24
    page.insert_text(fitz.Point(LEFT_MARGIN, y), designation_line, fontname=FONT_NAME_REGULAR, fontsize=13, color=COLOR_ACCENT)
    y += 22
    page.draw_line(fitz.Point(LEFT_MARGIN, y), fitz.Point(page.rect.width - RIGHT_MARGIN, y), color=COLOR_LINE, width=1)
    y += LINE_HEIGHT * 1.5
    
    for section in resume_data.get("sections", []):
        section_title = section.get("title", "Untitled").strip()
        title_lower = section_title.lower()

        if y > page.rect.height - BOTTOM_MARGIN - 80:
            page.insert_text(fitz.Point((page.rect.width - 10)/2, page.rect.height - BOTTOM_MARGIN/2), f"{page_num}", fontname=FONT_NAME_REGULAR, fontsize=9, color=COLOR_SECONDARY)
            page = doc.new_page(width=595, height=842)
            y = TOP_MARGIN
            page_num += 1
            if logo_path: page.insert_image(fitz.Rect(page.rect.width - RIGHT_MARGIN - 80, TOP_MARGIN - 40, page.rect.width - RIGHT_MARGIN, TOP_MARGIN - 10), filename=logo_path)
        
        y += LINE_HEIGHT * 1.8
        page.insert_text(fitz.Point(LEFT_MARGIN, y), section_title.upper(), fontname=FONT_NAME_BOLD, fontsize=14, color=COLOR_PRIMARY)
        y += 8
        page.draw_line(fitz.Point(LEFT_MARGIN, y), fitz.Point(page.rect.width - RIGHT_MARGIN, y), color=COLOR_ACCENT, width=0.5)
        y += LINE_HEIGHT
        
        content = section.get("content")
        
        if title_lower == "projects" and isinstance(content, list):
            for project_item in content:
                if not isinstance(project_item, dict): continue
                project_name = project_item.get("project_name", "N/A")
                description = project_item.get("description", "")
                tech_stack = project_item.get("tech_stack", "")
                rows_data = [("Project Name", project_name)]
                if description and str(description).strip(): rows_data.append(("Description", description))
                if tech_stack and str(tech_stack).strip(): rows_data.append(("Tech Stack", tech_stack))
                col1_width, col2_width = 120, page.rect.width - LEFT_MARGIN - RIGHT_MARGIN - 120
                table_height, row_heights = 0, []
                for label, data in rows_data:
                    h = _get_wrapped_text_height(str(data), col2_width - 10)
                    row_h = max(LINE_HEIGHT * 1.5, h) + 10
                    row_heights.append(row_h)
                    table_height += row_h
                if y + table_height > page.rect.height - BOTTOM_MARGIN:
                    page.insert_text(fitz.Point((page.rect.width - 10)/2, page.rect.height - BOTTOM_MARGIN/2), f"{page_num}", fontname=FONT_NAME_REGULAR, fontsize=9, color=COLOR_SECONDARY)
                    page = doc.new_page(width=595, height=842); y = TOP_MARGIN; page_num += 1
                    if logo_path: page.insert_image(fitz.Rect(page.rect.width - RIGHT_MARGIN - 80, TOP_MARGIN - 40, page.rect.width - RIGHT_MARGIN, TOP_MARGIN - 10), filename=logo_path)
                table_start_y = y
                for i, (label, data) in enumerate(rows_data):
                    row_h = row_heights[i]
                    page.draw_rect(fitz.Rect(LEFT_MARGIN, y, LEFT_MARGIN + col1_width, y + row_h), color=COLOR_TABLE_HEADER_BG, fill=COLOR_TABLE_HEADER_BG)
                    page.insert_textbox(fitz.Rect(LEFT_MARGIN + 5, y + 5, LEFT_MARGIN + col1_width - 5, y + row_h), label, fontsize=BODY_FONTSIZE, fontname=FONT_NAME_BOLD, color=COLOR_PRIMARY)
                    page.insert_textbox(fitz.Rect(LEFT_MARGIN + col1_width + 5, y + 5, page.rect.width - RIGHT_MARGIN - 5, y + row_h), str(data), fontsize=BODY_FONTSIZE, fontname=FONT_NAME_REGULAR, color=COLOR_SECONDARY)
                    y += row_h
                page.draw_rect(fitz.Rect(LEFT_MARGIN, table_start_y, page.rect.width - RIGHT_MARGIN, y), width=1, color=COLOR_LINE)
                page.draw_line(fitz.Point(LEFT_MARGIN + col1_width, table_start_y), fitz.Point(LEFT_MARGIN + col1_width, y), color=COLOR_LINE, width=1)
                temp_y = table_start_y
                for h in row_heights[:-1]:
                    temp_y += h
                    page.draw_line(fitz.Point(LEFT_MARGIN, temp_y), fitz.Point(page.rect.width - RIGHT_MARGIN, temp_y), color=COLOR_LINE, width=1)
                y += LINE_HEIGHT
        elif title_lower == "experience" and isinstance(content, list):
            for job in content:
                if not isinstance(job, dict): continue
                job_title = job.get('job_title', 'Untitled Job')
                company_date = job.get('company_and_date', '')
                duties = job.get('duties', [])
                
                est_height = LINE_HEIGHT * 2 + sum([_get_wrapped_text_height(d, page.rect.width - LEFT_MARGIN - RIGHT_MARGIN - 15) for d in duties])
                if y + est_height > page.rect.height - BOTTOM_MARGIN:
                     page.insert_text(fitz.Point((page.rect.width - 10)/2, page.rect.height - BOTTOM_MARGIN/2), f"{page_num}", fontname=FONT_NAME_REGULAR, fontsize=9, color=COLOR_SECONDARY)
                     page = doc.new_page(width=595, height=842); y = TOP_MARGIN; page_num += 1
                     if logo_path: page.insert_image(fitz.Rect(page.rect.width - RIGHT_MARGIN - 80, TOP_MARGIN - 40, page.rect.width - RIGHT_MARGIN, TOP_MARGIN - 10), filename=logo_path)

                page.insert_text(fitz.Point(LEFT_MARGIN, y), job_title, fontname=FONT_NAME_BOLD, fontsize=BODY_FONTSIZE + 1, color=COLOR_PRIMARY)
                y += LINE_HEIGHT
                if company_date:
                    page.insert_text(fitz.Point(LEFT_MARGIN, y), company_date, fontname=FONT_NAME_ITALIC, fontsize=BODY_FONTSIZE, color=COLOR_SECONDARY)
                    y += LINE_HEIGHT * 1.2

                for duty in duties:
                    duty_height = _get_wrapped_text_height(duty, page.rect.width - LEFT_MARGIN - RIGHT_MARGIN - 15)
                    if y + duty_height > page.rect.height - BOTTOM_MARGIN:
                        page.insert_text(fitz.Point((page.rect.width - 10)/2, page.rect.height - BOTTOM_MARGIN/2), f"{page_num}", fontname=FONT_NAME_REGULAR, fontsize=9, color=COLOR_SECONDARY)
                        page = doc.new_page(width=595, height=842); y = TOP_MARGIN; page_num += 1
                        if logo_path: page.insert_image(fitz.Rect(page.rect.width - RIGHT_MARGIN - 80, TOP_MARGIN - 40, page.rect.width - RIGHT_MARGIN, TOP_MARGIN - 10), filename=logo_path)
                    y = _render_text_with_wrapping(page, y, duty, page.rect.width - LEFT_MARGIN - RIGHT_MARGIN, x_offset=15, bullet=True)
                y += LINE_HEIGHT

        else: # Generic content (string or list of strings)
            content_list = content if isinstance(content, list) else [content] if content else []
            for item in content_list:
                item_height = _get_wrapped_text_height(item, page.rect.width - LEFT_MARGIN - RIGHT_MARGIN)
                if y + item_height > page.rect.height - BOTTOM_MARGIN:
                    page.insert_text(fitz.Point((page.rect.width - 10)/2, page.rect.height - BOTTOM_MARGIN/2), f"{page_num}", fontname=FONT_NAME_REGULAR, fontsize=9, color=COLOR_SECONDARY)
                    page = doc.new_page(width=595, height=842); y = TOP_MARGIN; page_num += 1
                    if logo_path: page.insert_image(fitz.Rect(page.rect.width - RIGHT_MARGIN - 80, TOP_MARGIN - 40, page.rect.width - RIGHT_MARGIN, TOP_MARGIN - 10), filename=logo_path)
                
                is_bullet = isinstance(content, list)
                y = _render_text_with_wrapping(page, y, str(item), page.rect.width - LEFT_MARGIN - RIGHT_MARGIN, x_offset=15 if is_bullet else 0, bullet=is_bullet)

    page.insert_text(fitz.Point((page.rect.width - 10)/2, page.rect.height - BOTTOM_MARGIN/2), f"{page_num}", fontname=FONT_NAME_REGULAR, fontsize=9, color=COLOR_SECONDARY)
    pdf_buffer = io.BytesIO()
    doc.save(pdf_buffer)
    doc.close()
    pdf_buffer.seek(0)
    return pdf_buffer

def create_docx(resume_data, company):
    """
    Generates a professional-looking DOCX resume from structured JSON data.
    """
    doc = Document()

    # Document-Wide Setup
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    list_bullet_style = doc.styles['List Bullet']
    list_bullet_format = list_bullet_style.paragraph_format
    list_bullet_format.left_indent = Inches(0.25)
    list_bullet_format.space_after = Pt(3)
    list_bullet_format.space_before = Pt(0)
    list_bullet_format.line_spacing = 1.05

    # Define Colors and Helper Functions
    TITLE_BLUE = RGBColor(47, 84, 150)
    SUBTLE_GRAY = RGBColor(89, 89, 89)

    def add_section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.font.name = 'Calibri'
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = TITLE_BLUE
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr.set(qn('w:bottom'), 'w:single w:sz="6" w:space="1" w:color="2F5496"')
        pPr.append(pBdr)

    # Logo in Header
    logo_paths = { "beround": "static/logos/beround.jpg", "climber": "static/logos/climber.jpg", "rennova": "static/logos/rennova.jpg" }
    logo_path = logo_paths.get(company)
    if logo_path and os.path.exists(logo_path):
        header = doc.sections[0].header
        htable = header.add_table(rows=1, cols=1, width=Inches(6.5))
        htable.alignment = WD_TABLE_ALIGNMENT.RIGHT
        htab_cell = htable.cell(0, 0)
        htab_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = htab_cell.paragraphs[0].add_run()
        run.add_picture(logo_path, height=Inches(0.5))

    # Main Header (Name, Title, Contact)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.columns[0].width = Inches(4.75)
    header_table.columns[1].width = Inches(2.25)

    cell_name = header_table.cell(0, 0)
    cell_name.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p_name = cell_name.paragraphs[0]
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run(resume_data.get("candidate_name", "Candidate Name"))
    run_name.font.name = 'Calibri Light'
    run_name.font.size = Pt(28)

    p_designation = cell_name.add_paragraph(resume_data.get("designation_line", "Professional Title"))
    p_designation.paragraph_format.space_before = Pt(0)
    for run in p_designation.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = SUBTLE_GRAY

    cell_contact = header_table.cell(0, 1)
    cell_contact.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    contact_info = resume_data.get("contact_info", {})
    for value in filter(None, contact_info.values()):
        p = cell_contact.add_paragraph(value)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    if len(cell_contact.paragraphs) > len(list(filter(None, contact_info.values()))):
        p_to_delete = cell_contact.paragraphs[0]._p
        p_to_delete.getparent().remove(p_to_delete)

    # Loop Through Resume Sections
    for section in resume_data.get("sections", []):
        section_title = section.get("title", "Untitled").strip()
        content = section.get("content")
        if not section_title or not content:
            continue
        
        add_section_title(section_title)
        
        title_lower = section_title.lower().strip()

        if title_lower == "experience" and isinstance(content, list):
            for job in content:
                if not isinstance(job, dict): continue
                p_job_title = doc.add_paragraph()
                p_job_title.paragraph_format.space_before = Pt(6)
                p_job_title.paragraph_format.space_after = Pt(0)
                p_job_title.add_run(job.get('job_title', 'Untitled Job')).bold = True

                p_company = doc.add_paragraph()
                p_company.paragraph_format.space_before = Pt(0)
                p_company.paragraph_format.space_after = Pt(4)
                run_company = p_company.add_run(job.get('company_and_date', ''))
                run_company.italic = True
                run_company.font.size = Pt(10)
                run_company.font.color.rgb = SUBTLE_GRAY

                duties = job.get('duties', [])
                if isinstance(duties, list):
                    for duty in duties:
                        if str(duty).strip():
                            doc.add_paragraph(str(duty), style='List Bullet')
        elif title_lower == "projects" and isinstance(content, list):
            for project in content:
                if not isinstance(project, dict): continue
                p_proj_title = doc.add_paragraph()
                p_proj_title.paragraph_format.space_before = Pt(6)
                p_proj_title.paragraph_format.space_after = Pt(2)
                p_proj_title.add_run(project.get('project_name', 'Untitled Project')).bold = True
                
                if project.get('description'):
                    doc.add_paragraph(str(project.get('description')), style='List Bullet')
                
                if project.get('tech_stack'):
                    p_tech = doc.add_paragraph(style='List Bullet')
                    p_tech.add_run("Technologies: ").bold = True
                    p_tech.add_run(str(project.get('tech_stack')))
        elif 'skill' in title_lower and isinstance(content, list):
            skills_table = doc.add_table(rows=len(content), cols=2)
            skills_table.autofit = False
            skills_table.allow_autofit = False
            skills_table.columns[0].width = Inches(1.5)
            skills_table.columns[1].width = Inches(5.5)
            
            for i, item in enumerate(content):
                category, skills = "", ""
                if isinstance(item, dict):
                    category = item.get("category", "")
                    skills = item.get("skills", "")
                elif isinstance(item, str) and ":" in item:
                    parts = item.split(":", 1)
                    category = parts[0].strip()
                    skills = parts[1].strip()
                else:
                    skills = str(item)

                cat_cell = skills_table.cell(i, 0)
                cat_run = cat_cell.paragraphs[0].add_run(category)
                cat_run.font.bold = True
                
                skills_cell = skills_table.cell(i, 1)
                skills_cell.text = skills
                
                for cell in [cat_cell, skills_cell]:
                    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
                    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        else:
            if isinstance(content, list):
                for item in content:
                    if str(item).strip():
                        doc.add_paragraph(str(item), style='List Bullet')
            else:
                if str(content).strip():
                    doc.add_paragraph(str(content))
    
    # Save to Buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer